import os
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from docx import Document
from docx.shared import Pt, RGBColor
import pandas as pd
import subprocess
import cloudinary
import cloudinary.uploader
from dotenv import load_dotenv
load_dotenv()

# Configure Cloudinary
cloudinary.config(
    cloud_name=os.getenv("CLOUD_NAME"),
    api_key=os.getenv("API_KEY"),
    api_secret=os.getenv("API_SECRET")
)

def upload_pdf_to_folder(file_path, folder_name):
    try:
        # Upload PDF to Cloudinary in the specified folder
        result = cloudinary.uploader.upload(
            file_path,
            folder=folder_name,
            use_filename=True,      # Use the original filename
            unique_filename=False   # Do not append random characters
        )
        print("Upload successful. File URL:", result["secure_url"])
    except Exception as e:
        print("Upload failed:", str(e))

# FastAPI app
app = FastAPI()

# Allow CORS for frontend interaction
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://127.0.0.1:5500"],  # Add your frontend URL here
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Utility functions
def resolve_path(*args):
    return os.path.abspath(os.path.join(*args))

def replace_placeholders(doc, placeholders, font_styles):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in placeholders.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            run = paragraph.add_run(replacement)
                            style = font_styles.get(placeholder, {})
                            run.font.size = style.get('size', Pt(12))
                            run.font.italic = style.get('italic', False)
                            run.font.bold = style.get('bold', False)
                            run.font.color.rgb = style.get('color', RGBColor(0, 0, 0))
                            run.font.name = style.get('font', 'Calibri')

def convert_docx_to_pdf(docx_path, name, email, event_name):
    pdf_path = os.path.join(os.path.dirname(docx_path), f"{name}_{email}_certificate.pdf")
    subprocess.run(f"convertword {docx_path} {pdf_path}", shell=True, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    os.remove(docx_path)
    upload_pdf_to_folder(pdf_path, f"AutoCred/{event_name}")
    return pdf_path

def generate_certificate(template_path, placeholders, output_folder, event_name):
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    font_styles = {
        '{name}': {'size': Pt(28), 'italic': True, 'color': RGBColor(171, 124, 52), 'font': 'Georgia'},
        '{event}': {'size': Pt(21.5), 'italic': True, 'color': RGBColor(171, 124, 52), 'font': 'Georgia'},
        '{date}': {'size': Pt(21.5), 'italic': True, 'bold': True, 'color': RGBColor(171, 124, 52), 'font': 'Palatino Linotype'},
        '{department}': {'size': Pt(21.5), 'italic': True, 'color': RGBColor(171, 124, 52), 'font': 'Georgia'},
        '{year}': {'size': Pt(21.5), 'italic': True, 'color': RGBColor(171, 124, 52), 'font': 'Georgia'},
    }

    replace_placeholders(doc, placeholders, font_styles)
    email = placeholders.get('{email}', 'unknown').replace(" ", "_").strip()
    name = placeholders['{name}'].replace(" ", "_").strip()
    file_name = f"{name}_{email}_certificate.docx"
    docx_path = os.path.join(output_folder, file_name)
    doc.save(docx_path)
    pdf_path = convert_docx_to_pdf(docx_path, name, email, event_name)
    return pdf_path

# API Endpoints
@app.post("/generate-certificates")
async def generate_certificates(
    event_name: str = Form(...),
    event_date: str = Form(...),
    template: str = Form(...),
    gen_type: str = Form(...),
    file: UploadFile = File(None),
    student_name: str = Form(None),
    department: str = Form(None),
    year: str = Form(None),
    email: str = Form(None),
):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_participation = resolve_path(script_dir, "temp1.docx")
    template_organizer = resolve_path(script_dir, "temp2.docx")
    pdf_folder = "/tmp/Certificates"
    os.makedirs(pdf_folder, exist_ok=True)
    template_path = template_participation if template == "template1" else template_organizer

    try:
        if gen_type == "bulk" and file:
            excel_path = resolve_path(script_dir, file.filename)
            with open(excel_path, "wb") as f:
                f.write(file.file.read())
            data = pd.read_excel(excel_path, engine="openpyxl").to_dict(orient="records")
            pdf_files = []
            for student in data:
                placeholders = {
                    '{name}': student['Name'],
                    '{department}': student['Department'],
                    '{year}': student['Year'],
                    '{event}': event_name,
                    '{date}': event_date,
                    '{email}': student.get('Email', 'unknown')
                }
                pdf_path = generate_certificate(template_path, placeholders, pdf_folder, placeholders['{event}'])
                pdf_files.append(pdf_path)
            return JSONResponse(content={"message": "Bulk certificates generated successfully.", "pdf_files": pdf_files})
        elif gen_type == "single":
            placeholders = {
                '{name}': student_name,
                '{department}': department,
                '{year}': year,
                '{event}': event_name,
                '{date}': event_date,
                '{email}': email
            }
            pdf_path = generate_certificate(template_path, placeholders, pdf_folder, placeholders['{event}'])
            return JSONResponse(content={"message": "Single certificate generated successfully.", "pdf_path": pdf_path})
        return JSONResponse(content={"error": "Invalid generation type or missing data"}, status_code=400)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
